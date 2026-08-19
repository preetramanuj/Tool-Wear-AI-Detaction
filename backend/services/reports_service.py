import io
import os
import csv
import time
import datetime
from typing import Dict, Any, List, Optional
from sqlalchemy.orm import Session
from sqlalchemy import func

from backend.core.config import settings
from backend.database.models import (
    Tool,
    InspectionRecord,
    Machine,
    AlertRecord,
    MaintenanceEvent,
    DowntimeEvent,
    EconomicParameters,
)
from backend.database.crud import get_economic_parameters
from backend.services.manufacturing_insights_service import manufacturing_insights_service
from backend.services.economic_impact_service import economic_impact_service
from backend.services.downtime_service import downtime_service

# ReportLab imports for PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
    HRFlowable,
)

# python-docx for Word DOCX generation
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

class ReportsService:
    """
    Industrial Report Generation Engine for ToolGuard-AI.
    Compiles live SQLite audit data into structured JSON, PDF documents,
    and Word (.docx) files for plant managers and quality auditors.
    """

    def generate_report_data(
        self,
        db: Session,
        report_type: str = "COMPREHENSIVE_AUDIT",
        timeframe: str = "ALL",
        machine_id: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Gathers comprehensive data across all system models and records.
        """
        now = datetime.datetime.now()
        report_id = f"RPT-{now.strftime('%Y%m%d')}-{int(time.time() * 1000) % 10000:04d}"

        # 1. Query Tools
        tools_query = db.query(Tool)
        if machine_id and machine_id != "ALL":
            tools_query = tools_query.filter(Tool.machine_id == machine_id)
        tools = tools_query.all()

        # 2. Query Inspections
        inspections_query = db.query(InspectionRecord).order_by(InspectionRecord.timestamp.desc())
        if machine_id and machine_id != "ALL":
            inspections_query = inspections_query.filter(InspectionRecord.machine_id == machine_id)
        
        # Apply Timeframe
        if timeframe == "24H":
            cutoff = now - datetime.timedelta(hours=24)
            inspections_query = inspections_query.filter(InspectionRecord.timestamp >= cutoff)
        elif timeframe == "7D":
            cutoff = now - datetime.timedelta(days=7)
            inspections_query = inspections_query.filter(InspectionRecord.timestamp >= cutoff)
        elif timeframe == "30D":
            cutoff = now - datetime.timedelta(days=30)
            inspections_query = inspections_query.filter(InspectionRecord.timestamp >= cutoff)
            
        inspections = inspections_query.limit(100).all()

        # 3. Query Machines & Alerts
        machines = db.query(Machine).all()
        alerts = db.query(AlertRecord).order_by(AlertRecord.timestamp.desc()).limit(20).all()

        # 4. Intelligence Services
        insights_data = manufacturing_insights_service.generate_comprehensive_insights(db)
        econ_data = economic_impact_service.calculate_economic_impact(db)
        downtime_data = downtime_service.get_downtime_analytics(db)

        # 5. Aggregate KPIs
        total_tools = len(tools)
        healthy_count = sum(1 for t in tools if t.status == "HEALTHY")
        warning_count = sum(1 for t in tools if t.status == "WARNING")
        critical_count = sum(1 for t in tools if t.status == "CRITICAL")
        retired_count = sum(1 for t in tools if t.status == "RETIRED")

        wears = [t.current_wear_um for t in tools if t.current_wear_um is not None and t.current_wear_um > 0]
        avg_wear_um = round(sum(wears) / len(wears), 1) if wears else 0.0

        ruls = [t.current_rul_cycles for t in tools if t.current_rul_cycles is not None]
        avg_rul = round(sum(ruls) / len(ruls), 0) if ruls else 0.0

        currency = econ_data.get("currency", "₹")

        # Titles and Subtitles
        titles = {
            "COMPREHENSIVE_AUDIT": "Facility Tool Wear & Machine Reliability Comprehensive Audit",
            "DAILY_INSPECTIONS": "Daily Tool Wear Inspection & Quality Verification Log",
            "DEGRADATION_TRENDS": "Tool Degradation Rates & Remaining Useful Life Analysis",
            "ECONOMIC_RELIABILITY": "Production Downtime Avoidance & Maintenance Cost Audit",
        }
        title = titles.get(report_type, "Industrial Tool Health & Compliance Audit Report")

        # Formatted Tools List
        tools_summary = []
        for t in tools:
            tools_summary.append({
                "tool_id": t.tool_id,
                "tool_name": t.tool_name,
                "tool_type": t.tool_type,
                "material": t.material,
                "coating": t.coating,
                "machine_id": t.machine_id,
                "current_wear_um": round(t.current_wear_um, 1),
                "current_wear_vb_mm": round(t.current_wear_vb_mm, 4),
                "current_rul_cycles": round(t.current_rul_cycles, 1) if t.current_rul_cycles is not None else None,
                "wear_rate": round(t.current_wear_rate, 3) if t.current_wear_rate is not None else None,
                "status": t.status,
                "total_inspections": t.total_inspections,
            })

        # Formatted Inspections List
        inspections_list = []
        for insp in inspections:
            inspections_list.append({
                "inspection_id": insp.inspection_id,
                "timestamp": insp.timestamp.strftime("%Y-%m-%d %H:%M") if insp.timestamp else "N/A",
                "tool_id": insp.tool_id,
                "machine_id": insp.machine_id,
                "operator_id": insp.operator_id,
                "wear_vb_mm": round(insp.wear_value, 3) if insp.wear_value is not None else 0.0,
                "wear_um": round(insp.wear_um, 1) if insp.wear_um is not None else 0.0,
                "health_status": insp.health_status,
                "rul_cycles": round(insp.rul_cycles, 1) if insp.rul_cycles is not None else None,
                "tool_eligibility": getattr(insp, "tool_eligibility", "ELIGIBLE"),
            })

        return {
            "report_id": report_id,
            "generated_at": now.strftime("%Y-%m-%d %H:%M:%S"),
            "generated_date": now.strftime("%B %d, %Y"),
            "report_type": report_type,
            "title": title,
            "report_title": title,
            "timeframe": timeframe,
            "machine_filter": machine_id or "ALL",
            "organization": "ToolGuard-AI Industrial Monitoring",
            "facility": "CNC Precision Manufacturing Cell #01-04",
            "executive_summary": insights_data.get("summary", "Plant tools evaluated within operational tolerances."),
            "summary_narrative": insights_data.get("summary", "Plant tools evaluated within operational tolerances."),
            "kpis": {
                "total_tools": total_tools,
                "healthy_tools": healthy_count,
                "warning_tools": warning_count,
                "critical_tools": critical_count,
                "retired_tools": retired_count,
                "total_inspections": len(inspections),
                "avg_wear_um": avg_wear_um,
                "avg_rul_cycles": avg_rul,
                "total_downtime_hours": downtime_data.get("summary", {}).get("total_downtime_hours", 0.0),
                "estimated_downtime_avoided_hours": downtime_data.get("summary", {}).get("estimated_downtime_avoided_hours", 0.0),
                "estimated_potential_savings": econ_data.get("summary", {}).get("estimated_potential_savings", {}).get("value", 0.0),
                "currency": currency,
            },
            "tools": tools_summary,
            "tools_summary": tools_summary,
            "inspections": inspections_list,
            "records": inspections_list,
            "maintenance_candidates": insights_data.get("maintenance_candidates", []),
            "machine_comparison": insights_data.get("machine_comparison", []),
            "insights": insights_data.get("insights", []),
            "economics": econ_data,
            "downtime": downtime_data,
        }

    def generate_pdf_report(self, report_data: Dict[str, Any]) -> bytes:
        """
        Renders a clean, executive PDF report using ReportLab.
        """
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = getSampleStyleSheet()
        
        # Custom Typography Styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#0F172A'),
        )
        subtitle_style = ParagraphStyle(
            'DocSubTitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=12,
            textColor=colors.HexColor('#64748B'),
        )
        h2_style = ParagraphStyle(
            'Heading2_Custom',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#0284C7'),
            spaceBefore=12,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            'Body_Custom',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=13,
            textColor=colors.HexColor('#334155'),
        )
        body_bold = ParagraphStyle(
            'Body_Bold',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=13,
            textColor=colors.HexColor('#0F172A'),
        )
        table_text = ParagraphStyle(
            'TableText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=8,
            leading=10,
            textColor=colors.HexColor('#1E293B'),
        )
        table_header = ParagraphStyle(
            'TableHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=8,
            leading=10,
            textColor=colors.HexColor('#0F172A'),
        )

        story = []

        # 1. Header Banner
        header_table_data = [
            [
                Paragraph(f"<b>TOOLGUARD-AI</b> | Precision Predictive Maintenance", ParagraphStyle('H1', parent=body_bold, textColor=colors.HexColor('#0284C7'), fontSize=10)),
                Paragraph(f"<b>Report ID:</b> {report_data['report_id']}", ParagraphStyle('H2', parent=body_style, alignment=2, fontSize=8)),
            ],
            [
                Paragraph(report_data['title'], title_style),
                Paragraph(f"<b>Date:</b> {report_data['generated_date']}<br/><b>Cell:</b> {report_data['facility']}", ParagraphStyle('H3', parent=body_style, alignment=2, fontSize=8)),
            ]
        ]
        header_table = Table(header_table_data, colWidths=[360, 160])
        header_table.setStyle(TableStyle([
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 8))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#0284C7'), spaceBefore=2, spaceAfter=10))

        # 2. Executive Summary Callout Box
        summary_text = f"<b>Facility Audit Summary:</b> {report_data['summary_narrative']}"
        callout_data = [[Paragraph(summary_text, body_style)]]
        callout_table = Table(callout_data, colWidths=[520])
        callout_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#F0F9FF')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#BAE6FD')),
            ('PADDING', (0,0), (-1,-1), 8),
        ]))
        story.append(callout_table)
        story.append(Spacer(1, 10))

        # 3. Key Operational Metrics Matrix
        k = report_data['kpis']
        curr = k.get('currency', '₹')
        kpi_data = [
            [
                Paragraph("<b>Monitored Tools</b>", table_header),
                Paragraph("<b>Active Warnings</b>", table_header),
                Paragraph("<b>Critical State</b>", table_header),
                Paragraph("<b>Average Wear</b>", table_header),
                Paragraph("<b>Avoided Downtime</b>", table_header),
                Paragraph("<b>Estimated Savings</b>", table_header),
            ],
            [
                Paragraph(str(k['total_tools']), ParagraphStyle('K1', parent=table_text, fontSize=11, fontName='Helvetica-Bold')),
                Paragraph(str(k['warning_tools']), ParagraphStyle('K2', parent=table_text, fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#D97706'))),
                Paragraph(str(k['critical_tools']), ParagraphStyle('K3', parent=table_text, fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#E11D48'))),
                Paragraph(f"{k['avg_wear_um']} µm", ParagraphStyle('K4', parent=table_text, fontSize=11, fontName='Helvetica-Bold')),
                Paragraph(f"{k['estimated_downtime_avoided_hours']} hrs", ParagraphStyle('K5', parent=table_text, fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#059669'))),
                Paragraph(f"{curr}{k['estimated_potential_savings']:,.0f}", ParagraphStyle('K6', parent=table_text, fontSize=11, fontName='Helvetica-Bold', textColor=colors.HexColor('#059669'))),
            ]
        ]
        kpi_table = Table(kpi_data, colWidths=[86, 86, 86, 86, 88, 88])
        kpi_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F8FAFC')),
            ('BACKGROUND', (0,1), (-1,1), colors.HexColor('#FFFFFF')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#E2E8F0')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
            ('PADDING', (0,0), (-1,-1), 6),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ]))
        story.append(kpi_table)
        story.append(Spacer(1, 12))

        # 4. Tool Inventory & Flank Wear Degradation Table
        story.append(Paragraph("1. Tool Insert Status & Degradation Registry", h2_style))
        tools_table_data = [
            [
                Paragraph("<b>Tool ID</b>", table_header),
                Paragraph("<b>Insert Type</b>", table_header),
                Paragraph("<b>Machine</b>", table_header),
                Paragraph("<b>Wear VB (mm)</b>", table_header),
                Paragraph("<b>Wear (µm)</b>", table_header),
                Paragraph("<b>RUL Cycles</b>", table_header),
                Paragraph("<b>Status</b>", table_header),
            ]
        ]
        for t in report_data['tools'][:15]: # Limit to top 15 for clean PDF fit
            status_color = '#059669' if t['status'] == 'HEALTHY' else ('#D97706' if t['status'] == 'WARNING' else '#E11D48')
            tools_table_data.append([
                Paragraph(t['tool_id'], table_text),
                Paragraph(t['tool_name'], table_text),
                Paragraph(t['machine_id'], table_text),
                Paragraph(f"{t['current_wear_vb_mm']:.3f}", table_text),
                Paragraph(f"{t['current_wear_um']:.1f}", table_text),
                Paragraph(str(t['current_rul_cycles']) if t['current_rul_cycles'] else "-", table_text),
                Paragraph(f"<b>{t['status']}</b>", ParagraphStyle('S1', parent=table_text, textColor=colors.HexColor(status_color))),
            ])
        
        tools_table = Table(tools_table_data, colWidths=[80, 140, 75, 75, 55, 45, 50])
        tools_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F8FAFC')]),
            ('PADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(tools_table)
        story.append(Spacer(1, 10))

        # 5. Recent Inspection Runs Table
        story.append(Paragraph("2. Optical Vision Inspection Log", h2_style))
        insp_table_data = [
            [
                Paragraph("<b>Inspection ID</b>", table_header),
                Paragraph("<b>Timestamp</b>", table_header),
                Paragraph("<b>Tool ID</b>", table_header),
                Paragraph("<b>Operator</b>", table_header),
                Paragraph("<b>Wear (VB)</b>", table_header),
                Paragraph("<b>Health</b>", table_header),
                Paragraph("<b>Eligibility</b>", table_header),
            ]
        ]
        for insp in report_data['inspections'][:10]:
            insp_table_data.append([
                Paragraph(insp['inspection_id'], table_text),
                Paragraph(insp['timestamp'], table_text),
                Paragraph(insp['tool_id'] or "-", table_text),
                Paragraph(insp['operator_id'] or "-", table_text),
                Paragraph(f"{insp['wear_vb_mm']:.3f} mm", table_text),
                Paragraph(insp['health_status'] or "-", table_text),
                Paragraph(insp.get('tool_eligibility', 'ELIGIBLE'), table_text),
            ])

        insp_table = Table(insp_table_data, colWidths=[90, 95, 85, 85, 65, 55, 45])
        insp_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#F1F5F9')),
            ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#CBD5E1')),
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#FFFFFF'), colors.HexColor('#F8FAFC')]),
            ('PADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(insp_table)
        story.append(Spacer(1, 10))

        # 6. Downtime & Economic Breakdown
        story.append(Paragraph("3. Production Reliability & Downtime Avoidance", h2_style))
        dt = report_data.get('downtime', {}).get('summary', {})
        rel_text = (
            f"During the audited monitoring interval, total recorded stoppage duration was <b>{dt.get('total_downtime_hours', 0.0):.1f} hours</b> "
            f"(Planned: {dt.get('planned_downtime_hours', 0.0):.1f}h, Unplanned: {dt.get('unplanned_downtime_hours', 0.0):.1f}h). "
            f"Through predictive tool condition pre-emption, an estimated <b>{dt.get('estimated_downtime_avoided_hours', 0.0):.1f} hours</b> of catastrophic spindle downtime were avoided, "
            f"delivering an estimated cost avoidance of <b>{curr}{dt.get('estimated_avoided_cost', 0.0):,.2f}</b>."
        )
        story.append(Paragraph(rel_text, body_style))
        story.append(Spacer(1, 14))

        # Footer Notice
        footer_text = "Generated automatically by ToolGuard-AI Vision & Predictive Intelligence System. Compliant with ISO standard cutting tool wear limits (300 µm VB EOL)."
        story.append(Paragraph(footer_text, ParagraphStyle('Foot', parent=subtitle_style, alignment=1, fontSize=7)))

        # Build PDF
        doc.build(story)
        pdf_bytes = buffer.getvalue()
        buffer.close()
        return pdf_bytes

    def generate_docx_report(self, report_data: Dict[str, Any]) -> bytes:
        """
        Generates a professionally styled Microsoft Word (.docx) document.
        """
        doc = docx.Document()
        
        # Set Page Margins (0.75 inch)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # 1. Document Title
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(report_data['title'])
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(15, 23, 42) # #0F172A

        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(f"Report ID: {report_data['report_id']} | Date: {report_data['generated_date']} | Cell: {report_data['facility']}")
        meta_run.font.size = Pt(9.5)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 2. Executive Summary Box
        sum_p = doc.add_paragraph()
        sum_run_b = sum_p.add_run("Automated Facility Summary: ")
        sum_run_b.bold = True
        sum_run_b.font.color.rgb = RGBColor(2, 132, 199) # #0284C7
        sum_p.add_run(report_data['summary_narrative'])
        sum_p.paragraph_format.space_after = Pt(12)

        # 3. KPI Summary Table
        k = report_data['kpis']
        curr = k.get('currency', '₹')
        kpi_table = doc.add_table(rows=2, cols=6)
        kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = ["Total Tools", "Healthy", "Warning", "Critical", "Avoided Downtime", "Estimated Savings"]
        vals = [
            str(k['total_tools']),
            str(k['healthy_tools']),
            str(k['warning_tools']),
            str(k['critical_tools']),
            f"{k['estimated_downtime_avoided_hours']} hrs",
            f"{curr}{k['estimated_potential_savings']:,.0f}"
        ]

        for col_idx, text in enumerate(headers):
            cell = kpi_table.cell(0, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for col_idx, text in enumerate(vals):
            cell = kpi_table.cell(1, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            if col_idx == 2:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(217, 119, 6)
            elif col_idx == 3:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(225, 29, 72)
            elif col_idx >= 4:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(5, 150, 105)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # 4. Tool Registry Table
        h2 = doc.add_heading("1. Tool Insert Degradation & Lifecycle Registry", level=2)
        h2.paragraph_format.space_before = Pt(12)
        h2.paragraph_format.space_after = Pt(6)

        tools_table = doc.add_table(rows=1, cols=7)
        tools_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        t_headers = ["Tool ID", "Insert Type", "Machine Cell", "Wear VB (mm)", "Wear (µm)", "RUL (Cyc)", "Status"]
        for col_idx, text in enumerate(t_headers):
            cell = tools_table.cell(0, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for tool in report_data['tools']:
            row_cells = tools_table.add_row().cells
            row_cells[0].text = tool['tool_id']
            row_cells[1].text = tool['tool_name']
            row_cells[2].text = tool['machine_id']
            row_cells[3].text = f"{tool['current_wear_vb_mm']:.3f}"
            row_cells[4].text = f"{tool['current_wear_um']:.1f}"
            row_cells[5].text = str(tool['current_rul_cycles']) if tool['current_rul_cycles'] else "-"
            row_cells[6].text = tool['status']
            for c in row_cells:
                c.paragraphs[0].runs[0].font.size = Pt(8.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(12)

        # 5. Inspection Log Table
        h3 = doc.add_heading("2. Recent Optical Inspection Records", level=2)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)

        insp_table = doc.add_table(rows=1, cols=6)
        insp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        i_headers = ["Inspection ID", "Timestamp", "Tool ID", "Operator", "Flank Wear VB", "Condition"]
        for col_idx, text in enumerate(i_headers):
            cell = insp_table.cell(0, col_idx)
            cell.text = text
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for insp in report_data['inspections'][:15]:
            row_cells = insp_table.add_row().cells
            row_cells[0].text = insp['inspection_id']
            row_cells[1].text = insp['timestamp']
            row_cells[2].text = insp['tool_id'] or "-"
            row_cells[3].text = insp['operator_id'] or "-"
            row_cells[4].text = f"{insp['wear_vb_mm']:.3f} mm"
            row_cells[5].text = insp['health_status'] or "-"
            for c in row_cells:
                c.paragraphs[0].runs[0].font.size = Pt(8.5)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        return docx_bytes

    def generate_csv_report(self, report_data: Dict[str, Any]) -> str:
        """
        Generates CSV format export string.
        """
        output = io.StringIO()
        writer = csv.writer(output)
        
        writer.writerow(["ToolGuard-AI Industrial Inspection Log"])
        writer.writerow(["Report ID", report_data["report_id"], "Generated Date", report_data["generated_date"]])
        writer.writerow([])
        
        # Tool Inventory
        writer.writerow(["--- TOOL INVENTORY ---"])
        writer.writerow(["Tool ID", "Name", "Type", "Material", "Coating", "Machine", "Wear (um)", "Wear VB (mm)", "RUL Cycles", "Status", "Inspections Count"])
        for t in report_data["tools"]:
            writer.writerow([
                t["tool_id"], t["tool_name"], t["tool_type"], t["material"], t["coating"],
                t["machine_id"], t["current_wear_um"], t["current_wear_vb_mm"],
                t["current_rul_cycles"] or "", t["status"], t["total_inspections"]
            ])
            
        writer.writerow([])
        # Inspection Logs
        writer.writerow(["--- INSPECTION RUNS ---"])
        writer.writerow(["Inspection ID", "Timestamp", "Tool ID", "Machine ID", "Operator ID", "Wear VB (mm)", "Wear (um)", "Health Status", "RUL Cycles", "Eligibility"])
        for insp in report_data["inspections"]:
            writer.writerow([
                insp["inspection_id"], insp["timestamp"], insp["tool_id"], insp["machine_id"],
                insp["operator_id"], insp["wear_vb_mm"], insp["wear_um"],
                insp["health_status"], insp["rul_cycles"] or "", insp.get("tool_eligibility", "ELIGIBLE")
            ])
            
        return output.getvalue()

reports_service = ReportsService()
